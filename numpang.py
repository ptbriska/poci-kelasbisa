import json
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from html import escape

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


# ============================================================
# CONFIG
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

# File Word sumber
INPUT_DOCX = BASE_DIR / "input" / "bank_soal.docx"

# Folder output
OUTPUT_DIR = BASE_DIR / "output"

# JSON database
JSON_FILE = OUTPUT_DIR / "data" / "questions.json"

# Asset gambar
ASSETS_DIR = OUTPUT_DIR / "assets" / "questions"

# Kalau True, gambar akan dikonversi ke WebP
CONVERT_IMAGES_TO_WEBP = True

# Kualitas WebP
WEBP_QUALITY = 85

# Maksimal lebar gambar.
# None = tidak resize.
# Contoh 1600 = gambar lebih besar dari 1600px akan diperkecil.
MAX_IMAGE_WIDTH = 1600

# Apakah otomatis melakukan git add/commit/push?
ENABLE_GIT_PUSH = False

# Commit message
GIT_COMMIT_MESSAGE = "Update bank soal"


# ============================================================
# FIELD WORD
# ============================================================

REQUIRED_FIELDS = [
    "ID SOAL",
    "SOAL",
    "OPSI A",
    "OPSI B",
    "OPSI C",
    "OPSI D",
    "KUNCI",
    "PEMBAHASAN",
]


# ============================================================
# UTIL
# ============================================================

def log(message=""):
    print(message)


def fail(message):
    print(f"\n❌ ERROR: {message}")
    sys.exit(1)


def normalize_label(text):
    """
    Normalisasi label seperti:
    'ID SOAL'
    'Id Soal'
    ' id soal '
    """
    text = text.strip().upper()
    text = re.sub(r"\s+", " ", text)
    return text


def sanitize_id(value):
    """
    ID soal hanya boleh:
    huruf
    angka
    underscore
    dash
    """
    value = value.strip()

    if not re.fullmatch(r"[A-Za-z0-9_-]+", value):
        raise ValueError(
            f"ID soal '{value}' tidak valid. "
            "Gunakan hanya huruf, angka, '-' atau '_'."
        )

    return value


def clean_text(text):
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def ensure_directories():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    JSON_FILE.parent.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# WORD XML HELPERS
# ============================================================

def get_relationship_target(document, rel_id):
    """
    Mengambil target file dari relationship ID.
    """
    try:
        rel = document.part.rels[rel_id]
        return rel.target_part
    except KeyError:
        return None


def get_image_rel_id(drawing_element):
    """
    Mendapatkan rId gambar dari elemen Drawing.
    """
    blip = drawing_element.find(
        ".//" + qn("a:blip")
    )

    if blip is None:
        return None

    return (
        blip.get(qn("r:embed"))
        or blip.get(qn("r:link"))
    )


def is_drawing(element):
    """
    Apakah XML element merupakan gambar/drawing?
    """
    return element.tag in {
        qn("w:drawing"),
        qn("w:pict"),
    }


# ============================================================
# IMAGE EXTRACTION
# ============================================================

class ImageManager:
    def __init__(self, document, question_id):
        self.document = document
        self.question_id = question_id

        self.output_dir = (
            ASSETS_DIR / question_id
        )

        self.output_dir.mkdir(
            parents=True,
            exist_ok=True
        )

        self.counter = 0
        self.cache = {}

    def extract(self, drawing):
        """
        Ekstrak gambar dari drawing Word.

        Return:
            relative asset path
        """

        rel_id = get_image_rel_id(drawing)

        if not rel_id:
            return None

        # Kalau gambar yang sama muncul berkali-kali,
        # gunakan file yang sama.
        if rel_id in self.cache:
            return self.cache[rel_id]

        part = get_relationship_target(
            self.document,
            rel_id
        )

        if part is None:
            return None

        self.counter += 1

        original_name = Path(
            part.partname
        ).name

        original_ext = (
            Path(original_name)
            .suffix
            .lower()
        )

        if CONVERT_IMAGES_TO_WEBP:

            filename = (
                f"image-{self.counter:02d}.webp"
            )

            output_file = (
                self.output_dir /
                filename
            )

            self.save_as_webp(
                part.blob,
                output_file
            )

        else:

            ext = original_ext

            if ext not in {
                ".png",
                ".jpg",
                ".jpeg",
                ".gif",
                ".webp",
            }:
                ext = ".png"

            filename = (
                f"image-{self.counter:02d}{ext}"
            )

            output_file = (
                self.output_dir /
                filename
            )

            output_file.write_bytes(
                part.blob
            )

        relative_path = (
            f"assets/questions/"
            f"{self.question_id}/"
            f"{filename}"
        )

        self.cache[rel_id] = relative_path

        return relative_path

    def save_as_webp(self, data, output_file):

        temp_file = (
            output_file.parent /
            f".tmp_{output_file.name}"
        )

        temp_file.write_bytes(data)

        try:

            with Image.open(temp_file) as img:

                # Handle transparency
                if img.mode in ("RGBA", "LA"):
                    converted = img.convert("RGBA")
                else:
                    converted = img.convert("RGB")

                # Resize jika terlalu besar
                if (
                    MAX_IMAGE_WIDTH
                    and converted.width > MAX_IMAGE_WIDTH
                ):
                    ratio = (
                        MAX_IMAGE_WIDTH /
                        converted.width
                    )

                    new_height = int(
                        converted.height * ratio
                    )

                    converted = converted.resize(
                        (
                            MAX_IMAGE_WIDTH,
                            new_height
                        ),
                        Image.Resampling.LANCZOS
                    )

                converted.save(
                    output_file,
                    "WEBP",
                    quality=WEBP_QUALITY,
                    method=6
                )

        except Exception as e:

            # Jika gagal konversi, simpan file asli
            shutil.copy2(
                temp_file,
                output_file
            )

            log(
                f"   ⚠ Gagal convert gambar: {e}"
            )

        finally:

            if temp_file.exists():
                temp_file.unlink()


# ============================================================
# HTML PARSER
# ============================================================

def paragraph_to_html(
    paragraph,
    document,
    image_manager
):
    """
    Mengubah satu paragraph Word menjadi HTML.

    Mempertahankan:
    - teks
    - bold
    - italic
    - underline
    - gambar
    """

    parts = []

    p = paragraph._p

    for child in p.iterchildren():

        # ----------------------------------------------------
        # RUN
        # ----------------------------------------------------

        if child.tag == qn("w:r"):

            text_parts = []

            for node in child.iterchildren():

                if node.tag == qn("w:t"):

                    text = node.text or ""

                    text_parts.append(
                        escape(text)
                    )

                elif node.tag == qn("w:tab"):

                    text_parts.append(
                        "&emsp;"
                    )

                elif node.tag == qn("w:br"):

                    text_parts.append(
                        "<br>"
                    )

                elif node.tag == qn("w:drawing"):

                    image_path = (
                        image_manager.extract(
                            node
                        )
                    )

                    if image_path:

                        parts.append(
                            f'<img src="{image_path}" '
                            f'class="question-image" '
                            f'loading="lazy">'
                        )

            text = "".join(text_parts)

            if text:

                rpr = child.find(
                    qn("w:rPr")
                )

                bold = False
                italic = False
                underline = False

                if rpr is not None:

                    bold = (
                        rpr.find(qn("w:b"))
                        is not None
                    )

                    italic = (
                        rpr.find(qn("w:i"))
                        is not None
                    )

                    underline = (
                        rpr.find(qn("w:u"))
                        is not None
                    )

                if bold:
                    text = f"<strong>{text}</strong>"

                if italic:
                    text = f"<em>{text}</em>"

                if underline:
                    text = f"<u>{text}</u>"

                parts.append(text)

        # ----------------------------------------------------
        # DRAWING LANGSUNG
        # ----------------------------------------------------

        elif child.tag == qn("w:drawing"):

            image_path = (
                image_manager.extract(
                    child
                )
            )

            if image_path:

                parts.append(
                    f'<img src="{image_path}" '
                    f'class="question-image" '
                    f'loading="lazy">'
                )

    content = "".join(parts).strip()

    if not content:
        return ""

    return f"<p>{content}</p>"


def cell_to_html(
    cell,
    document,
    image_manager
):
    """
    Mengubah seluruh isi cell menjadi HTML.

    Gambar akan tetap mengikuti urutan
    teks di Word.
    """

    html_parts = []

    for paragraph in cell.paragraphs:

        html = paragraph_to_html(
            paragraph,
            document,
            image_manager
        )

        if html:
            html_parts.append(html)

    return "\n".join(html_parts)


# ============================================================
# TABLE PARSER
# ============================================================

def parse_question_table(
    table,
    document,
    table_number
):
    """
    Membaca satu tabel Word menjadi satu soal.

    Format:
    kolom 1 = label
    kolom 2 = isi
    """

    fields = {}

    for row in table.rows:

        if len(row.cells) < 2:
            continue

        label = normalize_label(
            row.cells[0].text
        )

        if not label:
            continue

        fields[label] = row.cells[1]

    # Cek field wajib
    missing = [
        field
        for field in REQUIRED_FIELDS
        if field not in fields
    ]

    if missing:

        raise ValueError(
            f"Tabel #{table_number}: "
            f"field tidak ditemukan: "
            f"{', '.join(missing)}"
        )

    # ID
    question_id = clean_text(
        fields["ID SOAL"].text
    )

    question_id = sanitize_id(
        question_id
    )

    image_manager = ImageManager(
        document,
        question_id
    )

    # --------------------------------------------------------
    # QUESTION
    # --------------------------------------------------------

    question_html = cell_to_html(
        fields["SOAL"],
        document,
        image_manager
    )

    # --------------------------------------------------------
    # OPTIONS
    # --------------------------------------------------------

    options = {}

    for letter in ["A", "B", "C", "D"]:

        options[letter] = cell_to_html(
            fields[f"OPSI {letter}"],
            document,
            image_manager
        )

    # --------------------------------------------------------
    # ANSWER
    # --------------------------------------------------------

    answer = clean_text(
        fields["KUNCI"].text
    ).upper()

    if answer not in ["A", "B", "C", "D"]:
        raise ValueError(
            f"{question_id}: "
            f"KUNCI harus A/B/C/D, "
            f"ditemukan '{answer}'"
        )

    # --------------------------------------------------------
    # EXPLANATION
    # --------------------------------------------------------

    explanation_html = cell_to_html(
        fields["PEMBAHASAN"],
        document,
        image_manager
    )

    # --------------------------------------------------------
    # RESULT
    # --------------------------------------------------------

    return {
        "id": question_id,

        "question": {
            "html": question_html
        },

        "options": options,

        "answer": answer,

        "explanation": {
            "html": explanation_html
        }
    }


# ============================================================
# VALIDATION
# ============================================================

def validate_question(
    question,
    existing_ids
):
    errors = []

    question_id = question["id"]

    # Duplicate
    if question_id in existing_ids:

        errors.append(
            f"ID duplikat: {question_id}"
        )

    # Question kosong
    if not question["question"]["html"].strip():

        errors.append(
            f"{question_id}: soal kosong."
        )

    # Option kosong
    for letter, content in (
        question["options"].items()
    ):

        if not content.strip():

            errors.append(
                f"{question_id}: "
                f"opsi {letter} kosong."
            )

    # Explanation
    if not question["explanation"]["html"].strip():

        log(
            f"   ⚠ {question_id}: "
            "pembahasan kosong."
        )

    return errors


# ============================================================
# JSON
# ============================================================

def save_json(questions):

    data = {
        "version": 2,
        "generated_at": None,
        "questions": questions
    }

    with open(
        JSON_FILE,
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            data,
            file,
            ensure_ascii=False,
            indent=2
        )


# ============================================================
# GIT
# ============================================================

def run_git(command):

    result = subprocess.run(
        command,
        cwd=BASE_DIR,
        capture_output=True,
        text=True
    )

    if result.returncode != 0:

        raise RuntimeError(
            result.stderr.strip()
        )

    return result.stdout.strip()


def publish_to_git():

    log("\n================================")
    log("GIT PUBLISH")
    log("================================")

    try:

        log("git add...")

        run_git([
            "git",
            "add",
            "output/data/questions.json",
            "output/assets/"
        ])

        # Cek apakah ada perubahan
        status = run_git([
            "git",
            "status",
            "--porcelain"
        ])

        if not status:

            log("Tidak ada perubahan.")
            return

        log("git commit...")

        run_git([
            "git",
            "commit",
            "-m",
            GIT_COMMIT_MESSAGE
        ])

        log("git push...")

        run_git([
            "git",
            "push"
        ])

        log("✓ GitHub berhasil diperbarui.")

    except Exception as e:

        log(
            f"❌ Git publish gagal: {e}"
        )


# ============================================================
# MAIN
# ============================================================

def main():

    print()
    print("==============================================")
    print("        CBT WORD QUESTION IMPORTER")
    print("==============================================")
    print()

    # --------------------------------------------------------
    # CHECK INPUT
    # --------------------------------------------------------

    if not INPUT_DOCX.exists():

        fail(
            f"File Word tidak ditemukan:\n"
            f"{INPUT_DOCX}"
        )

    ensure_directories()

    # --------------------------------------------------------
    # LOAD WORD
    # --------------------------------------------------------

    log("📄 Membaca Word...")
    log(f"   {INPUT_DOCX}")

    try:

        document = Document(
            INPUT_DOCX
        )

    except Exception as e:

        fail(
            f"Gagal membuka Word: {e}"
        )

    tables = document.tables

    log(
        f"✓ Ditemukan {len(tables)} tabel."
    )

    if not tables:

        fail(
            "Tidak ditemukan tabel soal."
        )

    # --------------------------------------------------------
    # PARSE
    # --------------------------------------------------------

    questions = []
    ids = set()

    log()
    log("🔍 Memproses soal...")

    for index, table in enumerate(
        tables,
        start=1
    ):

        log(
            f"\n[{index}/{len(tables)}] "
            f"Memproses tabel..."
        )

        try:

            question = parse_question_table(
                table,
                document,
                index
            )

            errors = validate_question(
                question,
                ids
            )

            if errors:

                for error in errors:
                    log(f"   ❌ {error}")

                continue

            questions.append(
                question
            )

            ids.add(
                question["id"]
            )

            log(
                f"   ✓ {question['id']}"
            )

        except Exception as e:

            log(
                f"   ❌ GAGAL: {e}"
            )

    # --------------------------------------------------------
    # RESULT
    # --------------------------------------------------------

    log()
    log("==============================================")
    log("HASIL IMPORT")
    log("==============================================")

    log(
        f"Soal berhasil : {len(questions)}"
    )

    log(
        f"Soal gagal    : "
        f"{len(tables) - len(questions)}"
    )

    if not questions:

        fail(
            "Tidak ada soal yang berhasil diimport."
        )

    # --------------------------------------------------------
    # SAVE JSON
    # --------------------------------------------------------

    log()
    log("💾 Membuat questions.json...")

    save_json(
        questions
    )

    log(
        f"✓ {JSON_FILE}"
    )

    # --------------------------------------------------------
    # GIT
    # --------------------------------------------------------

    if ENABLE_GIT_PUSH:

        publish_to_git()

    # --------------------------------------------------------
    # DONE
    # --------------------------------------------------------

    log()
    log("==============================================")
    log("✅ SELESAI")
    log("==============================================")
    log()

    log(
        f"JSON   : {JSON_FILE}"
    )

    log(
        f"ASSETS : {ASSETS_DIR}"
    )

    log()


if __name__ == "__main__":
    main()
